"""
Text extraction for the ILM Generator.

Supports: PDF, DOCX, TXT, EPUB, JPG, JPEG, PNG.
DJVU is intentionally not supported - there is no reliable pure-Python
DJVU library, and the djvulibre system binary is not dependably available
on Streamlit Cloud. Convert DJVU files to PDF before uploading.

Scanned (image-only) PDF pages are automatically OCR'd as a fallback
when the page's real text layer is empty or nearly empty.

Embedded diagram images (a labeled figure sitting inside an otherwise
text-rich page/document - e.g. a CPU block diagram, an apparatus setup
drawing) are ALSO detected and OCR'd separately, even when the page
already has plenty of its own text. Without this, a diagram whose labels
exist only inside the image (not in the document's real text layer) is
completely invisible to Gemini, and gets silently dropped from the
generated material. Each detected diagram becomes a
"[DIAGRAM ON THIS PAGE: ...]" (PDF) or "[DIAGRAM IN DOCUMENT: ...]"
(DOCX) marker in the extracted text, carrying whatever OCR could read
from the image.
"""

import io
import zipfile

import fitz  # PyMuPDF
from docx import Document

from ocr import ocr_image_bytes, ocr_pdf_page

SUPPORTED_EXTENSIONS = {"pdf", "docx", "txt", "epub", "jpg", "jpeg", "png"}

# Roughly how many characters make up one printed page of study material.
# Used only to estimate a "page count" for non-PDF formats, where there is
# no real concept of a page (DOCX/TXT/EPUB) or where 1 file = 1 page (images).
_CHARS_PER_PAGE_ESTIMATE = 1800

# A PDF page is treated as "scanned" (no real text layer) if it has fewer
# than this many extracted characters, triggering an OCR fallback for it.
_MIN_CHARS_BEFORE_OCR_FALLBACK = 20

# An embedded image smaller than this (in pixels, either dimension) is
# treated as decorative (a bullet, icon, logo, divider line) and skipped -
# it is very unlikely to be a genuine diagram/figure.
_MIN_DIAGRAM_IMAGE_DIMENSION_PX = 80

# An embedded DOCX image smaller than this (in raw bytes) is treated as
# decorative and skipped, mirroring the pixel-dimension check used for
# PDF images above (DOCX media files don't expose pixel dimensions
# without fully decoding the image, so byte size is used as a cheap proxy).
_MIN_DOCX_IMAGE_BYTES = 2000

# OCR text shorter than this from a genuine-sized image is treated as
# "no readable labels" rather than a failed/empty diagram - the marker
# still notes that an image is present, just without OCR content.
_MIN_DIAGRAM_OCR_CHARS = 8


def get_extension(filename):
    return filename.rsplit(".", 1)[-1].lower() if "." in filename else ""


def is_supported(filename):
    return get_extension(filename) in SUPPORTED_EXTENSIONS


# -----------------------------------------------------------------
# Shared helper: turn one embedded image's bytes into a diagram marker
# -----------------------------------------------------------------

def _diagram_marker_for_image(image_bytes, label="DIAGRAM ON THIS PAGE"):
    """
    OCRs a single embedded image and returns a "[LABEL: ...]" marker
    string, or None if the image is too small/OCR failed outright (as
    opposed to succeeding with little/no text, which still returns a
    marker noting an image is present).
    """
    try:
        ocr_text = (ocr_image_bytes(image_bytes) or "").strip()
    except Exception:
        return None

    if len(ocr_text) >= _MIN_DIAGRAM_OCR_CHARS:
        return f"[{label}: {ocr_text}]"
    return (
        f"[{label}: an embedded figure/diagram is present here; "
        "see the surrounding text for what it likely shows]"
    )


# -----------------------------------------------------------------
# PDF
# -----------------------------------------------------------------

def _extract_embedded_image_diagrams(document, page):
    """
    Finds meaningfully-sized embedded images on a PDF page and OCRs each
    one, returning a list of "[DIAGRAM ON THIS PAGE: ...]" marker
    strings. Runs regardless of whether the page already has a normal
    text layer, since a diagram's own labels are not part of that text
    layer at all.
    """
    markers = []
    try:
        images = page.get_images(full=True)
    except Exception:
        return markers

    seen_xrefs = set()
    for img in images:
        xref = img[0]
        if xref in seen_xrefs:
            continue  # the same image can be listed more than once per page
        seen_xrefs.add(xref)

        try:
            base_image = document.extract_image(xref)
        except Exception:
            continue

        width = base_image.get("width", 0)
        height = base_image.get("height", 0)
        if width < _MIN_DIAGRAM_IMAGE_DIMENSION_PX or height < _MIN_DIAGRAM_IMAGE_DIMENSION_PX:
            continue  # too small to plausibly be a diagram (icon/bullet/rule)

        image_bytes = base_image.get("image")
        if not image_bytes:
            continue

        marker = _diagram_marker_for_image(image_bytes, label="DIAGRAM ON THIS PAGE")
        if marker:
            markers.append(marker)

    return markers


def extract_pdf_text(file_bytes):
    """
    Extract text from a PDF, one page at a time. Any page with little or
    no real text (a scanned page) is rendered to an image and OCR'd
    automatically. Separately, every embedded image on every page is
    checked and OCR'd if it looks like a genuine diagram (see
    _extract_embedded_image_diagrams), so diagrams are captured even on
    pages that already have plenty of their own text.
    Returns (text, page_count).
    """
    document = fitz.open(stream=file_bytes, filetype="pdf")
    text = ""
    page_count = len(document)

    for page_number in range(page_count):
        page = document.load_page(page_number)
        page_text = page.get_text().strip()

        if len(page_text) < _MIN_CHARS_BEFORE_OCR_FALLBACK:
            ocr_text = ocr_pdf_page(page)
            if ocr_text:
                page_text = ocr_text

        diagram_markers = _extract_embedded_image_diagrams(document, page)
        if diagram_markers:
            page_text = (page_text + "\n\n" + "\n\n".join(diagram_markers)).strip()

        text += f"\n\n===== PAGE {page_number + 1} =====\n\n"
        text += page_text

    document.close()
    return text, page_count


# -----------------------------------------------------------------
# DOCX
# -----------------------------------------------------------------

def _extract_docx_image_diagrams(file_bytes):
    """
    Extracts every embedded image from a .docx package (word/media/*)
    and OCRs each one, returning "[DIAGRAM IN DOCUMENT: ...]" marker
    strings. DOCX has no real concept of "pages", so these markers are
    appended after the main extracted text rather than interleaved at an
    exact position - what matters is that the diagram's content is
    visible to Gemini at all, not its precise placement.
    """
    markers = []
    try:
        with zipfile.ZipFile(io.BytesIO(file_bytes)) as z:
            media_files = sorted(n for n in z.namelist() if n.startswith("word/media/"))
            for name in media_files:
                try:
                    image_bytes = z.read(name)
                except Exception:
                    continue
                if len(image_bytes) < _MIN_DOCX_IMAGE_BYTES:
                    continue  # too small to plausibly be a diagram (icon/bullet/logo)

                marker = _diagram_marker_for_image(image_bytes, label="DIAGRAM IN DOCUMENT")
                if marker:
                    markers.append(marker)
    except Exception:
        pass
    return markers


def extract_docx_text(file_bytes):
    document = Document(io.BytesIO(file_bytes))
    paragraphs = [p.text for p in document.paragraphs if p.text.strip()]

    for table in document.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                paragraphs.append(row_text)

    text = "\n".join(paragraphs)

    diagram_markers = _extract_docx_image_diagrams(file_bytes)
    if diagram_markers:
        text = (text + "\n\n" + "\n\n".join(diagram_markers)).strip()

    page_count = max(1, len(text) // _CHARS_PER_PAGE_ESTIMATE)
    return text, page_count


# -----------------------------------------------------------------
# TXT
# -----------------------------------------------------------------

def extract_txt_text(file_bytes):
    for encoding in ("utf-8", "utf-16", "latin-1"):
        try:
            text = file_bytes.decode(encoding)
            break
        except UnicodeDecodeError:
            continue
    else:
        text = file_bytes.decode("utf-8", errors="ignore")

    page_count = max(1, len(text) // _CHARS_PER_PAGE_ESTIMATE)
    return text, page_count


# -----------------------------------------------------------------
# EPUB
# -----------------------------------------------------------------

def extract_epub_text(file_bytes):
    import ebooklib
    from ebooklib import epub
    from bs4 import BeautifulSoup

    tmp_path = "/tmp/_ilm_upload.epub"
    with open(tmp_path, "wb") as f:
        f.write(file_bytes)

    book = epub.read_epub(tmp_path)
    chunks = []
    for item in book.get_items():
        if item.get_type() == ebooklib.ITEM_DOCUMENT:
            soup = BeautifulSoup(item.get_content(), "html.parser")
            chunk_text = soup.get_text(separator="\n").strip()
            if chunk_text:
                chunks.append(chunk_text)

    text = "\n\n".join(chunks)
    page_count = max(1, len(text) // _CHARS_PER_PAGE_ESTIMATE)
    return text, page_count


# -----------------------------------------------------------------
# IMAGES (jpg / jpeg / png) - OCR only
# -----------------------------------------------------------------

def extract_image_text(file_bytes):
    text = ocr_image_bytes(file_bytes)
    return text, 1  # one image counts as one "page"


# -----------------------------------------------------------------
# DISPATCH
# -----------------------------------------------------------------

def extract_text(file_bytes, filename):
    """
    Extracts text from a single file of any supported type.
    Returns (text, estimated_page_count).
    Raises ValueError for unsupported file types (including .djvu).
    """
    ext = get_extension(filename)

    if ext == "pdf":
        return extract_pdf_text(file_bytes)
    elif ext == "docx":
        return extract_docx_text(file_bytes)
    elif ext == "txt":
        return extract_txt_text(file_bytes)
    elif ext == "epub":
        return extract_epub_text(file_bytes)
    elif ext in ("jpg", "jpeg", "png"):
        return extract_image_text(file_bytes)
    elif ext == "djvu":
        raise ValueError(
            "DJVU files are not supported. Please convert this file to PDF "
            "and upload it again."
        )
    else:
        raise ValueError(f"Unsupported file type: .{ext}")


def extract_multiple(files):
    """
    Extracts and combines text from a list of uploaded files.
    `files` is a list of (filename, file_bytes) tuples.

    Returns a dict:
        {
            "combined_text": str,
            "total_pages": int,
            "per_file": [{"filename": ..., "pages": ..., "chars": ...}, ...],
            "errors": [{"filename": ..., "error": ...}, ...],
        }
    A single failing file does not stop extraction of the others.
    """
    combined_parts = []
    per_file = []
    errors = []
    total_pages = 0

    for filename, file_bytes in files:
        try:
            text, pages = extract_text(file_bytes, filename)
            combined_parts.append(f"\n\n########## SOURCE FILE: {filename} ##########\n\n{text}")
            per_file.append({"filename": filename, "pages": pages, "chars": len(text)})
            total_pages += pages
        except Exception as e:
            errors.append({"filename": filename, "error": str(e)})

    return {
        "combined_text": "\n".join(combined_parts),
        "total_pages": total_pages,
        "per_file": per_file,
        "errors": errors,
    }
